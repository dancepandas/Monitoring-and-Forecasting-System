import argparse
import asyncio
import json
import logging
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[3]))

from gateway.config import settings
from gateway.services import data_cache

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

_DEVICE_CODE = "FD000489923695"


async def _get_cached(station_code: str):
    """从缓存读取水位和流量原始数据。"""
    level = await data_cache.get(f"aiflow:level:{station_code}", max_age=600)
    flow_raw = await data_cache.get(f"aiflow:flow_raw:{station_code}:{_DEVICE_CODE}", max_age=600)
    return level, flow_raw


async def generate_report(report_type, station_code, date_str):
    reports_dir = Path(settings.reports_dir)
    reports_dir.mkdir(parents=True, exist_ok=True)

    # 从缓存读取数据
    level_data, flow_data = await _get_cached(station_code)
    level_items = (level_data.get("data", []) or []) if level_data else []
    flow_raw_items = (flow_data.get("data", []) or []) if flow_data else []

    # 提取水位统计（优先来自 level 缓存，fallback 到 flow_raw 中的 waterLevel）
    wl_values = []
    for item in level_items:
        wl = item.get("waterLevel")
        if wl is not None:
            try: wl_values.append(float(wl))
            except: pass
    if not wl_values:
        for item in flow_raw_items:
            wl = item.get("waterLevel")
            if wl is not None:
                try: wl_values.append(float(wl))
                except: pass

    # 提取流量统计（从原始数据）
    flow_values = []
    for item in flow_raw_items:
        vf = item.get("virtualFlow")
        if vf is not None:
            try:
                flow_values.append(round(float(vf), 2))
            except (ValueError, TypeError):
                pass

    # 生成 Word
    ext = "docx" if HAS_DOCX else "txt"
    filename = f"{station_code}_{report_type}_{date_str}.{ext}"
    filepath = reports_dir / filename

    if HAS_DOCX:
        doc = Document()
        title = doc.add_heading(f"水文监测报告 - {station_code}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"报告类型: {report_type}")
        doc.add_paragraph(f"测站编码: {station_code}")
        doc.add_paragraph(f"日期: {date_str}")
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 水位汇总
        doc.add_heading("水位数据", level=1)
        doc.add_paragraph(f"共 {len(level_items)} 条记录")
        if wl_values:
            doc.add_paragraph(f"最高水位: {max(wl_values):.2f} m")
            doc.add_paragraph(f"最低水位: {min(wl_values):.2f} m")
            doc.add_paragraph(f"平均水位: {sum(wl_values)/len(wl_values):.2f} m")
        if level_items:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "水位(m)"
            hdr[2].text = "状态"
            for item in level_items[:50]:
                row = table.add_row().cells
                row[0].text = str(item.get("measureTime", ""))
                row[1].text = str(item.get("waterLevel", ""))
                row[2].text = str(item.get("status", ""))

        # 流量汇总
        doc.add_heading("流量数据", level=1)
        doc.add_paragraph(f"共 {len(flow_raw_items)} 条原始记录")
        if flow_values:
            doc.add_paragraph(f"最大流量: {max(flow_values):.0f} m³/s")
            doc.add_paragraph(f"最小流量: {min(flow_values):.0f} m³/s")
            doc.add_paragraph(f"平均流量: {sum(flow_values)/len(flow_values):.0f} m³/s")
            doc.add_paragraph(f"最新流量: {flow_values[-1]:.0f} m³/s (时间: {flow_raw_items[-1].get('measureTime', '')})")
        if flow_raw_items:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "流量(m³/s)"
            hdr[2].text = "水位(m)"
            for item in flow_raw_items[:50]:
                row = table.add_row().cells
                row[0].text = str(item.get("measureTime", ""))
                row[1].text = str(item.get("virtualFlow", ""))
                row[2].text = str(item.get("waterLevel", ""))

        doc.save(filepath)
    else:
        filepath = reports_dir / f"{station_code}_{report_type}_{date_str}.txt"
        content = f"水文监测报告 - {station_code}\n"
        content += f"报告类型: {report_type}\n日期: {date_str}\n\n"
        content += f"水位记录: {len(level_items)} 条\n"
        if wl_values:
            content += f"最高: {max(wl_values):.2f}m  最低: {min(wl_values):.2f}m  平均: {sum(wl_values)/len(wl_values):.2f}m\n"
        content += f"\n流量原始记录: {len(flow_raw_items)} 条\n"
        if flow_values:
            content += f"最大: {max(flow_values):.0f}m³/s  最小: {min(flow_values):.0f}m³/s  平均: {sum(flow_values)/len(flow_values):.0f}m³/s\n"
        filepath.write_text(content, encoding="utf-8")
        filename = filepath.name

    size = filepath.stat().st_size
    summary = f"{date_str} 水位 {len(level_items)} 条，流量 {len(flow_raw_items)} 条"
    if wl_values:
        summary += f"，最高水位 {max(wl_values):.2f}m"
    if flow_values:
        summary += f"，最新流量 {flow_values[-1]:.0f}m³/s"

    return {"path": str(filepath), "filename": filename, "size": size, "summary": summary}


def main():
    parser = argparse.ArgumentParser(description="Generate hydrology report")
    parser.add_argument("--type", default="daily", choices=["daily", "weekly", "monthly"])
    parser.add_argument("--station", required=True, help="Station code")
    parser.add_argument("--date", default=datetime.now().strftime("%Y-%m-%d"), help="Date YYYY-MM-DD")
    args = parser.parse_args()

    result = asyncio.run(generate_report(args.type, args.station, args.date))
    sys.stdout.reconfigure(encoding='utf-8')
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
